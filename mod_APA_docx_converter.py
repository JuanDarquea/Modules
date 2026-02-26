from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from tkinter import Tk, filedialog

# Ocultar la ventana principal de tkinter
root = Tk()
root.withdraw()

# Abrir diálogo para seleccionar archivo
print("📂 Selecciona el archivo DOCX a formatear...")
input_path = filedialog.askopenfilename(
    title="Seleccionar archivo DOCX",
    filetypes=[("Documentos Word", "*.docx"), ("Todos los archivos", "*.*")]
)

# Verificar si se seleccionó un archivo
if not input_path:
    print("❌ No se seleccionó ningún archivo. Operación cancelada.")
    exit()

# Generar nombre de salida automáticamente
file_dir = os.path.dirname(input_path)
file_name = os.path.basename(input_path)
file_name_without_ext = os.path.splitext(file_name)[0]
output_path = os.path.join(file_dir, f"{file_name_without_ext}_APA.docx")

try:
    print(f"📄 Procesando: {file_name}")
    
    # Cargar el documento
    doc = Document(input_path)
    
    # --- CONFIGURACIÓN GENERAL APA ---
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # --- FORMATO DE PÁRRAFOS ---
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        paragraph.paragraph_format.line_spacing = 2.0  # APA usa doble espaciado
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # --- ENCABEZADO APA ---
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        
        # Limpiar encabezado previo
        for p in header.paragraphs:
            p.clear()
        
        # Crear tabla invisible para alinear título (izq.) y página (der.)
        table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        
        # Título en la celda izquierda
        left_cell = table.rows[0].cells[0]
        left_paragraph = left_cell.paragraphs[0]
        left_paragraph.text = "IMPLEMENTACIÓN DE BIG DATA EN LA COOPERATIVA JEP"
        left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in left_paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        
        # Número de página en la celda derecha
        right_cell = table.rows[0].cells[1]
        right_paragraph = right_cell.paragraphs[0]
        right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run = right_paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = " PAGE "
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        
        # Hacer la tabla invisible (sin bordes)
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
    
    # Guardar el documento
    doc.save(output_path)
    print(f"\n✅ Documento formateado con éxito!")
    print(f"📁 Archivo original: {input_path}")
    print(f"💾 Archivo guardado: {output_path}")
    print(f"📍 Ubicación: {os.path.abspath(output_path)}")

except Exception as e:
    print(f"\n❌ Error al procesar el documento: {e}")
    import traceback
    traceback.print_exc()